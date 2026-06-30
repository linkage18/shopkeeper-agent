"""Convert interview-100-questions.md to a formatted .docx file."""
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def md_to_docx(md_path, docx_path):
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    
    # Set heading styles
    for level in range(1, 5):
        try:
            heading_style = doc.styles[f'Heading {level}']
            heading_style.font.name = '微软雅黑'
            heading_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            if level == 1:
                heading_style.font.size = Pt(22)
                heading_style.font.bold = True
            elif level == 2:
                heading_style.font.size = Pt(18)
                heading_style.font.bold = True
            elif level == 3:
                heading_style.font.size = Pt(14)
                heading_style.font.bold = True
            else:
                heading_style.font.size = Pt(12)
                heading_style.font.bold = True
        except:
            pass
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Split by lines
    lines = content.split('\n')
    
    # Track if we're in a table, code block, or blockquote
    in_code_block = False
    in_table = False
    in_blockquote = False
    
    for line in lines:
        stripped = line.strip()
        
        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                in_code_block = False
                p = doc.add_paragraph()
                continue
            else:
                in_code_block = True
                p = doc.add_paragraph()
                continue
        
        if in_code_block:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue
        
        # Handle empty lines
        if not stripped:
            # Skip consecutive empty lines
            continue
        
        # Handle horizontal rules
        if stripped == '---':
            p = doc.add_paragraph()
            p.space_before = Pt(6)
            p.space_after = Pt(6)
            # Add a bottom border to simulate HR
            p.style = doc.styles['Normal']
            pf = p.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(12)
            continue
        
        # Handle headings
        if stripped.startswith('#'):
            # Count heading level
            level = 0
            for ch in stripped:
                if ch == '#':
                    level += 1
                else:
                    break
            text = stripped[level:].strip()
            if text:
                if level <= 4:
                    doc.add_heading(text, level=level)
                else:
                    doc.add_heading(text, level=4)
            continue
        
        # Handle blockquotes
        if stripped.startswith('>'):
            text = stripped.lstrip('>').strip()
            if text:
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                pf = p.paragraph_format
                pf.left_indent = Inches(0.3)
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                run.font.italic = True
            continue
        
        # Handle unordered list items
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:].strip()
            if text:
                p = doc.add_paragraph(style='List Bullet')
                p.text = text
            continue
        
        # Handle ordered list items
        if re.match(r'^\d+[\.\)]\s', stripped):
            text = re.sub(r'^\d+[\.\)]\s', '', stripped)
            if text:
                p = doc.add_paragraph(style='List Number')
                p.text = text
            continue
        
        # Handle bold markers in inline text
        def process_inline(text):
            # Process **bold**
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # Process `code`
            text = re.sub(r'`([^`]+)`', r'\1', text)
            return text
        
        # Regular paragraph
        if stripped:
            text = process_inline(stripped)
            
            # Check if it starts with bold question pattern like "### 1. **Title**"
            bold_match = re.match(r'\*\*(.*?)\*\*(.*)', text)
            
            p = doc.add_paragraph()
            
            # Handle **考察点:** and **参考:** patterns
            if '**' in line and not text.startswith('**'):
                # Split by bold markers to create runs
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        inner = part[2:-2]
                        run = p.add_run(inner)
                        run.bold = True
                    elif part.strip():
                        p.add_run(part)
            else:
                p.add_run(text)
    
    # Add page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    doc.save(docx_path)
    print(f"✅ Saved: {docx_path}")

if __name__ == '__main__':
    md_to_docx('docs/interview-100-questions.md', 'docs/interview-100-questions.docx')
