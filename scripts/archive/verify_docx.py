from docx import Document
doc = Document("docs/interview-100-questions.docx")
paras = len(doc.paragraphs)
headings = len([p for p in doc.paragraphs if p.style.name.startswith("Heading")])
print(f"Paragraphs: {paras}, Headings: {headings}")
